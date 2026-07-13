from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import CondPageBreak, PageBreak, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "windows-setup" / "installer-assets" / "docs" / "USER_GUIDE.md"
OUTPUT_DIR = ROOT / "windows-setup" / "Installers"
DOCX_PATH = OUTPUT_DIR / "NexaFlow_User_Guide.docx"
PDF_PATH = OUTPUT_DIR / "NexaFlow_User_Guide.pdf"

BLUE = RGBColor(0x00, 0x96, 0xB7)
DARK = RGBColor(0x15, 0x23, 0x38)
MUTED = RGBColor(0x5E, 0x6B, 0x7A)


def parse_markdown(text: str):
    blocks = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            blocks.append(("space", ""))
        elif line.startswith("### "):
            blocks.append(("h3", line[4:]))
        elif line.startswith("## "):
            blocks.append(("h2", line[3:]))
        elif line.startswith("# "):
            blocks.append(("title", line[2:]))
        elif re.match(r"^\d+\. ", line):
            blocks.append(("number", re.sub(r"^\d+\. ", "", line)))
        elif line.startswith("- "):
            blocks.append(("bullet", line[2:]))
        else:
            blocks.append(("body", line))
    return blocks


def add_inline_runs(paragraph, text: str, *, color=DARK, size=11, bold=False):
    for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not part:
            continue
        run = paragraph.add_run(part.strip("*`") if part.startswith(("**", "`")) else part)
        run.font.name = "Calibri"
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.bold = bold or part.startswith("**")
        if part.startswith("`"):
            run.font.name = "Consolas"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")


def build_docx(blocks):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, DARK),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "NEXAFLOW  |  USER GUIDE"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.name = "Calibri"
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_inline_runs(footer, "NexaFlow v1.0.0  |  ", color=MUTED, size=8.5)
    page_field = OxmlElement("w:fldSimple")
    page_field.set(qn("w:instr"), "PAGE")
    footer._p.append(page_field)

    title_seen = False
    for kind, text in blocks:
        if kind == "space":
            continue
        if kind == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(116)
            p.paragraph_format.space_after = Pt(12)
            add_inline_runs(p, text, color=DARK, size=30, bold=True)
            title_seen = True
        elif title_seen and kind == "body" and len(doc.paragraphs) <= 3:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            add_inline_runs(p, text, color=MUTED, size=13)
            if text.startswith("Version"):
                doc.add_page_break()
        elif kind == "h2":
            if text.startswith("15. Support"):
                doc.add_page_break()
            p = doc.add_paragraph(style="Heading 1")
            add_inline_runs(p, text, color=BLUE, size=16, bold=True)
        elif kind == "h3":
            p = doc.add_paragraph(style="Heading 2")
            add_inline_runs(p, text, color=BLUE, size=13, bold=True)
        elif kind in ("bullet", "number"):
            style = "List Bullet" if kind == "bullet" else "List Number"
            p = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.188)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            add_inline_runs(p, text)
        else:
            p = doc.add_paragraph()
            add_inline_runs(p, text)

    core = doc.core_properties
    core.title = "NexaFlow v1 User Guide"
    core.subject = "Windows desktop automation and Android companion guide"
    core.author = "NexaFlow"
    core.last_modified_by = "NexaFlow"
    core.keywords = "NexaFlow, user guide, Windows, Android, automation"
    core.comments = ""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def pdf_text(text: str) -> str:
    text = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"`([^`]+)`", r"<font name='Courier'>\1</font>", text)
    return text.replace("&", "&amp;") if "&" in text else text


def build_pdf(blocks):
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=12.8, textColor=HexColor("#152326"), spaceAfter=5)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=15, leading=18, textColor=HexColor("#0096B7"), spaceBefore=12, spaceAfter=7)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=12.5, leading=15, textColor=HexColor("#0096B7"), spaceBefore=9, spaceAfter=5)
    bullet = ParagraphStyle("Bullet", parent=body, leftIndent=18, firstLineIndent=-9, bulletIndent=7, spaceAfter=4)
    title = ParagraphStyle("Title", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=29, leading=34, alignment=TA_CENTER, textColor=HexColor("#152326"), spaceAfter=12)
    subtitle = ParagraphStyle("Subtitle", parent=body, fontSize=13, leading=17, alignment=TA_CENTER, textColor=HexColor("#5E6B7A"), spaceAfter=12)

    story = [Spacer(1, 1.55 * inch)]
    title_body_count = 0
    number = 0
    for kind, text in blocks:
        if kind == "space":
            continue
        if kind == "title":
            story.append(Paragraph(pdf_text(text), title))
        elif kind == "body" and title_body_count < 2:
            story.append(Paragraph(pdf_text(text), subtitle))
            title_body_count += 1
            if title_body_count == 2:
                story.append(PageBreak())
        elif kind == "h2":
            number = 0
            if text.startswith("15. Support"):
                story.append(PageBreak())
            else:
                story.append(CondPageBreak(0.65 * inch))
            story.append(Paragraph(pdf_text(text), h1))
        elif kind == "h3":
            number = 0
            story.append(CondPageBreak(0.48 * inch))
            story.append(Paragraph(pdf_text(text), h2))
        elif kind == "bullet":
            number = 0
            story.append(Paragraph(pdf_text(text), bullet, bulletText="-"))
        elif kind == "number":
            number += 1
            story.append(Paragraph(pdf_text(text), bullet, bulletText=f"{number}."))
        else:
            number = 0
            story.append(Paragraph(pdf_text(text), body))

    def page(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica-Bold", 8)
        canvas.setFillColor(HexColor("#5E6B7A"))
        canvas.drawString(inch, 10.45 * inch, "NEXAFLOW  |  USER GUIDE")
        canvas.setFont("Helvetica", 8)
        canvas.drawRightString(7.5 * inch, 0.52 * inch, f"NexaFlow v1.0.0  |  {doc.page}")
        canvas.restoreState()

    document = SimpleDocTemplate(str(PDF_PATH), pagesize=LETTER, rightMargin=inch, leftMargin=inch, topMargin=inch, bottomMargin=inch, title="NexaFlow v1 User Guide", author="NexaFlow")
    document.build(story, onFirstPage=page, onLaterPages=page)


def main():
    blocks = parse_markdown(SOURCE.read_text(encoding="utf-8"))
    build_docx(blocks)
    build_pdf(blocks)
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    sys.exit(main())
